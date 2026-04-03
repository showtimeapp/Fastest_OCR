"""
DOC/DOCX Handler — extract text from Word documents (no OCR needed).

Uses:
  - python-docx for .docx files (preserves structure)
  - mammoth for .doc files (converts to markdown)
"""

from __future__ import annotations

import logging
from pathlib import Path

from src.ocr.postprocessor import (
    DocumentResult,
    PageResult,
    build_document_result,
)

logger = logging.getLogger(__name__)


def _extract_docx(file_path: Path) -> str:
    """Extract text from .docx using python-docx."""
    from docx import Document

    doc = Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Preserve heading levels
        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if rows:
            # Add markdown table header separator
            header = rows[0]
            num_cols = header.count("|") - 1
            separator = "| " + " | ".join(["---"] * num_cols) + " |"
            parts.append(rows[0])
            parts.append(separator)
            parts.extend(rows[1:])

    return "\n\n".join(parts)


def _extract_doc(file_path: Path) -> str:
    """Extract text from .doc using mammoth (converts to markdown)."""
    import mammoth

    with open(file_path, "rb") as f:
        result = mammoth.convert_to_markdown(f)

    if result.messages:
        for msg in result.messages:
            logger.debug("mammoth: %s", msg)

    return result.value


def process_doc_file(file_path: str | Path) -> DocumentResult:
    """
    Process a .doc or .docx file.
    No OCR involved — direct text extraction.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        content = _extract_docx(file_path)
    elif suffix == ".doc":
        content = _extract_doc(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    page_results = [
        PageResult(
            page_number=1,
            raw_text=content,
            regions=[
                {
                    "index": 0,
                    "label": "text",
                    "content": content,
                }
            ],
        )
    ]

    doc_result = build_document_result(
        filename=file_path.name,
        page_results=page_results,
        processing_time=0.001,
        file_type=suffix.lstrip("."),
    )

    logger.info("DOC file processed: %s (%d chars)", file_path.name, len(content))
    return doc_result
